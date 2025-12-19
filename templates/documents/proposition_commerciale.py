from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from pathlib import Path

class PropositionGenerator:
    def __init__(self):
        self.pricing = {"CH": {"RAG System": 25000, "Multi-Agent": 40000, "Training": 8000},
                        "DZ": {"RAG System": 800000, "Multi-Agent": 1500000, "Training": 300000}}
        self.contacts = {"CH": "contact@iafactory.ch", "DZ": "contact@iafactoryalgeria.com"}
    
    def generate(self, client: dict, services: list, market: str = "CH") -> str:
        doc = Document()
        currency = "CHF" if market == "CH" else "DZD"
        
        doc.add_heading("IA FACTORY", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_heading("PROPOSITION COMMERCIALE", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Client: {client.get('company', 'N/A')}")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%d/%m/%Y')}")
        doc.add_page_break()
        
        doc.add_heading("Services", 1)
        total = 0
        for s in services:
            if s in self.pricing[market]:
                doc.add_paragraph(f"• {s}: {self.pricing[market][s]:,} {currency}")
                total += self.pricing[market][s]
        doc.add_paragraph(f"TOTAL: {total:,} {currency}").runs[0].bold = True
        
        doc.add_heading("Contact", 1)
        doc.add_paragraph(f"Email: {self.contacts[market]}")
        
        Path("outputs/propositions").mkdir(parents=True, exist_ok=True)
        f = f"outputs/propositions/prop_{client.get('company','x').replace(' ','_')}_{market}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        doc.save(f)
        return f
