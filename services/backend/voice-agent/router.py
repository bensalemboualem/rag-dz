"""
FastAPI Router pour l'agent vocal Faster-Whisper
Endpoints API pour reconnaissance vocale souveraine
"""

from fastapi import APIRouter, File, UploadFile, Form, HTTPException
from fastapi.responses import JSONResponse
from typing import Optional
import logging

try:
    from .transcription_service import get_transcription_service
except ImportError:
    from transcription_service import get_transcription_service

logger = logging.getLogger(__name__)

# Router FastAPI
router = APIRouter(
    prefix="/api/voice-agent",
    tags=["voice-agent"],
)


@router.post("/transcribe")
async def transcribe_audio(
    file: UploadFile = File(..., description="Fichier audio à transcrire"),
    language: Optional[str] = Form(None, description="Code langue (fr, en, ar) ou auto"),
    professional_context: Optional[str] = Form(None, description="Contexte: medical, legal, accounting"),
):
    """
    Transcrit un fichier audio en texte

    **Use Cases Professionnels**:
    - **Médecins**: Comptes-rendus de consultation
    - **Avocats**: Notes d'audience, dictées juridiques
    - **Experts-comptables**: Notes de rendez-vous client

    **Langues supportées**:
    - `fr` - Français (France, Suisse, Belgique, Québec)
    - `en` - Anglais (US, UK, médical)
    - `ar` - Arabe (littéraire, dialectes, darija)
    - `null` - Détection automatique

    **Formats supportés** (audio + vidéo):
    - Audio: WAV, MP3, M4A, FLAC, OGG, OPUS, WEBM, AAC
    - Vidéo: MP4, AVI, MOV, MKV (l'audio est extrait automatiquement)

    **Réponse**:
    ```json
    {
      "text": "Texte complet transcrit",
      "cleaned_text": "Texte nettoyé selon contexte",
      "segments": [
        {"start": 0.0, "end": 2.5, "text": "Segment 1"},
        {"start": 2.5, "end": 5.0, "text": "Segment 2"}
      ],
      "language": "fr",
      "language_probability": 0.98,
      "duration": 45.3,
      "filename": "consultation_20250116.m4a",
      "professional_context": "medical"
    }
    ```

    **Exemple cURL**:
    ```bash
    curl -X POST "http://localhost:3000/api/voice-agent/transcribe" \\
      -F "file=@consultation.m4a" \\
      -F "language=fr" \\
      -F "professional_context=medical"
    ```
    """
    try:
        # Vérifier type de fichier (audio + vidéo - Whisper extrait l'audio automatiquement)
        allowed_extensions = [".wav", ".mp3", ".m4a", ".flac", ".ogg", ".opus", ".webm", ".aac", ".mp4", ".avi", ".mov", ".mkv"]
        file_ext = "." + file.filename.split(".")[-1].lower()

        if file_ext not in allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"Format non supporté: {file_ext}. Formats acceptés: {', '.join(allowed_extensions)}",
            )

        # Service de transcription
        service = get_transcription_service()

        # Transcription
        result = service.transcribe_file(
            audio_file=file.file,
            filename=file.filename,
            language=language,
            professional_context=professional_context,
        )

        logger.info(f"Transcription réussie: {file.filename} ({result['language']})")

        return JSONResponse(content=result)

    except Exception as e:
        logger.error(f"Erreur API transcribe: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/transcribe-url")
async def transcribe_from_url(
    audio_url: str = Form(..., description="URL du fichier audio"),
    language: Optional[str] = Form(None, description="Code langue ou auto"),
):
    """
    Transcrit un fichier audio depuis une URL

    **Exemple**:
    ```bash
    curl -X POST "http://localhost:3000/api/voice-agent/transcribe-url" \\
      -F "audio_url=https://example.com/audio.m4a" \\
      -F "language=fr"
    ```
    """
    try:
        service = get_transcription_service()
        result = service.transcribe_url(
            audio_url=audio_url,
            language=language,
        )

        return JSONResponse(content=result)

    except Exception as e:
        logger.error(f"Erreur transcribe-url: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/detect-language")
async def detect_language(
    file: UploadFile = File(..., description="Fichier audio"),
):
    """
    Détecte la langue d'un fichier audio

    **Réponse**:
    ```json
    {
      "language": "fr",
      "probability": 0.98
    }
    ```
    """
    try:
        service = get_transcription_service()
        result = service.detect_language(
            audio_file=file.file,
            filename=file.filename,
        )

        return JSONResponse(content=result)

    except Exception as e:
        logger.error(f"Erreur detect-language: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/models")
async def list_models():
    """
    Liste les modèles Whisper disponibles

    **Réponse**:
    ```json
    {
      "models": {
        "tiny": "Plus petit, plus rapide (39M params)",
        "base": "Modèle de base (74M params)",
        "small": "Petit modèle (244M params)",
        "medium": "Modèle moyen (769M params)",
        "large-v2": "Grand modèle v2 (1550M params)",
        "large-v3": "Grand modèle v3 - Recommandé (1550M params)",
        "distil-large-v3": "Version légère de large-v3 (50% plus rapide)"
      },
      "current_model": "large-v3",
      "device": "cuda",
      "compute_type": "float16"
    }
    ```
    """
    try:
        from .whisper_engine import WhisperEngine
    except ImportError:
        from whisper_engine import WhisperEngine

    return JSONResponse(
        content={
            "models": {
                "tiny": "Plus petit, plus rapide (39M params)",
                "base": "Modèle de base (74M params)",
                "small": "Petit modèle (244M params)",
                "medium": "Modèle moyen (769M params)",
                "large-v2": "Grand modèle v2 (1550M params)",
                "large-v3": "Grand modèle v3 - Recommandé (1550M params)",
                "distil-large-v3": "Version légère de large-v3 (50% plus rapide)",
            },
            "current_model": "large-v3",
            "device": "auto",
            "compute_type": "float16",
            "languages": [
                "fr (Français)",
                "en (English)",
                "ar (العربية)",
                "es (Español)",
                "de (Deutsch)",
                "it (Italiano)",
                "pt (Português)",
                "... 97 langues au total",
            ],
        }
    )


@router.get("/health")
async def health_check():
    """
    Health check de l'agent vocal

    **Réponse**:
    ```json
    {
      "status": "healthy",
      "service": "voice-agent",
      "model": "large-v3",
      "device": "cuda",
      "ready": true
    }
    ```
    """
    try:
        service = get_transcription_service()
        return JSONResponse(
            content={
                "status": "healthy",
                "service": "voice-agent",
                "model": "large-v3",
                "device": service.engine.device,
                "ready": True,
            }
        )
    except Exception as e:
        return JSONResponse(
            status_code=503,
            content={
                "status": "unhealthy",
                "service": "voice-agent",
                "error": str(e),
                "ready": False,
            },
        )


@router.post("/export-pdf")
async def export_to_pdf(
    text: str = Form(..., description="Texte à exporter"),
    title: str = Form("Transcription", description="Titre du document"),
    context: Optional[str] = Form(None, description="Contexte: medical, legal, accounting"),
):
    """
    Exporte une transcription en PDF

    **Exemple**:
    ```bash
    curl -X POST "http://localhost:3000/api/voice-agent/export-pdf" \\
      -F "text=Le patient présente..." \\
      -F "title=Consultation Dr. Martin" \\
      -F "context=medical"
    ```
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import cm
        from io import BytesIO
        import datetime

        # Créer PDF en mémoire
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                                rightMargin=2*cm, leftMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)

        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor='#10b981',
            spaceAfter=30,
        )

        # Contenu
        story = []
        story.append(Paragraph(title, title_style))
        story.append(Paragraph(f"Date: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}", styles['Normal']))
        if context:
            context_names = {"medical": "Médical", "legal": "Juridique", "accounting": "Comptable"}
            story.append(Paragraph(f"Contexte: {context_names.get(context, context)}", styles['Normal']))

        story.append(Spacer(1, 20))
        story.append(Paragraph(text.replace('\n', '<br/>'), styles['BodyText']))

        # Générer PDF
        doc.build(story)

        # Retourner le PDF
        buffer.seek(0)
        from fastapi.responses import StreamingResponse

        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={title.replace(' ', '_')}.pdf"}
        )

    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="reportlab not installed. Install with: pip install reportlab"
        )
    except Exception as e:
        logger.error(f"Erreur export PDF: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/export-docx")
async def export_to_docx(
    text: str = Form(..., description="Texte à exporter"),
    title: str = Form("Transcription", description="Titre du document"),
    context: Optional[str] = Form(None, description="Contexte: medical, legal, accounting"),
):
    """
    Exporte une transcription en DOCX (Word)

    **Exemple**:
    ```bash
    curl -X POST "http://localhost:3000/api/voice-agent/export-docx" \\
      -F "text=Le patient présente..." \\
      -F "title=Consultation Dr. Martin" \\
      -F "context=medical"
    ```
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from io import BytesIO
        import datetime

        # Créer document
        doc = Document()

        # Titre
        heading = doc.add_heading(title, level=1)
        heading.runs[0].font.color.rgb = RGBColor(16, 185, 129)

        # Métadonnées
        doc.add_paragraph(f"Date: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}")
        if context:
            context_names = {"medical": "Médical", "legal": "Juridique", "accounting": "Comptable"}
            doc.add_paragraph(f"Contexte: {context_names.get(context, context)}")

        doc.add_paragraph()  # Ligne vide

        # Contenu
        for paragraph in text.split('\n\n'):
            if paragraph.strip():
                p = doc.add_paragraph(paragraph)
                p.style = 'Body Text'

        # Sauvegarder en mémoire
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        from fastapi.responses import StreamingResponse

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={title.replace(' ', '_')}.docx"}
        )

    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx not installed. Install with: pip install python-docx"
        )
    except Exception as e:
        logger.error(f"Erreur export DOCX: {e}")
        raise HTTPException(status_code=500, detail=str(e))
