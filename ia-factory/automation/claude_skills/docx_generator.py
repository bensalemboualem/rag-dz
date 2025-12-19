"""
IA Factory - Claude Skill: DOCX Generator
Génération de documents Word professionnels
"""

from fastapi import APIRouter, HTTPException
from pydantic import BaseModel, Field
from typing import Optional, List, Dict, Any
from datetime import datetime
from enum import Enum
import os

router = APIRouter(prefix="/skills/docx", tags=["Claude Skills - DOCX"])


class DocumentType(str, Enum):
    PROPOSAL = "proposal"           # Proposition commerciale
    CONTRACT = "contract"           # Contrat
    REPORT = "report"               # Rapport
    INVOICE = "invoice"             # Facture
    LETTER = "letter"               # Lettre officielle
    MEMO = "memo"                   # Note interne
    SPECIFICATION = "specification" # Cahier des charges
    MANUAL = "manual"               # Manuel utilisateur
    CASE_STUDY = "case_study"       # Étude de cas


class DocumentStyle(str, Enum):
    PROFESSIONAL = "professional"
    MODERN = "modern"
    CLASSIC = "classic"
    MINIMALIST = "minimalist"
    CORPORATE = "corporate"


class DocumentRequest(BaseModel):
    """Requête de génération de document"""
    document_type: DocumentType
    title: str
    subtitle: Optional[str] = None
    author: str = "Boualem Chebaki"
    company: str = "IA Factory"
    recipient: Optional[str] = None
    recipient_company: Optional[str] = None
    sections: List[Dict[str, Any]] = Field(default_factory=list)
    style: DocumentStyle = DocumentStyle.PROFESSIONAL
    language: str = "fr"
    include_header: bool = True
    include_footer: bool = True
    include_toc: bool = False


class GeneratedDocument(BaseModel):
    """Document généré"""
    id: str
    filename: str
    filepath: str
    document_type: DocumentType
    created_at: datetime
    size_bytes: int
    download_url: str


class DocxGenerator:
    """
    Générateur de documents DOCX
    Utilise python-docx pour créer des documents professionnels
    """
    
    # Styles de couleurs par thème
    COLORS = {
        DocumentStyle.PROFESSIONAL: {
            "primary": "1F4E79",    # Bleu foncé
            "secondary": "2E75B6",  # Bleu moyen
            "accent": "BDD7EE",     # Bleu clair
            "text": "333333"        # Gris foncé
        },
        DocumentStyle.MODERN: {
            "primary": "00B050",    # Vert
            "secondary": "92D050",  # Vert clair
            "accent": "E2EFDA",     # Vert très clair
            "text": "262626"
        },
        DocumentStyle.CORPORATE: {
            "primary": "C00000",    # Rouge foncé
            "secondary": "FF6B6B",  # Rouge clair
            "accent": "FBE4D5",     # Beige
            "text": "404040"
        },
        DocumentStyle.MINIMALIST: {
            "primary": "000000",
            "secondary": "666666",
            "accent": "F2F2F2",
            "text": "000000"
        }
    }
    
    # Templates par type de document
    TEMPLATES = {
        DocumentType.PROPOSAL: {
            "sections": [
                {"name": "Résumé Exécutif", "required": True},
                {"name": "Contexte et Besoins", "required": True},
                {"name": "Solution Proposée", "required": True},
                {"name": "Méthodologie", "required": False},
                {"name": "Planning", "required": True},
                {"name": "Investissement", "required": True},
                {"name": "Pourquoi Nous Choisir", "required": False},
                {"name": "Prochaines Étapes", "required": True}
            ]
        },
        DocumentType.CONTRACT: {
            "sections": [
                {"name": "Parties", "required": True},
                {"name": "Objet du Contrat", "required": True},
                {"name": "Durée", "required": True},
                {"name": "Obligations des Parties", "required": True},
                {"name": "Conditions Financières", "required": True},
                {"name": "Confidentialité", "required": True},
                {"name": "Résiliation", "required": True},
                {"name": "Litiges", "required": True},
                {"name": "Signatures", "required": True}
            ]
        },
        DocumentType.REPORT: {
            "sections": [
                {"name": "Résumé", "required": True},
                {"name": "Introduction", "required": True},
                {"name": "Méthodologie", "required": False},
                {"name": "Résultats", "required": True},
                {"name": "Analyse", "required": True},
                {"name": "Recommandations", "required": True},
                {"name": "Conclusion", "required": True},
                {"name": "Annexes", "required": False}
            ]
        },
        DocumentType.SPECIFICATION: {
            "sections": [
                {"name": "Introduction", "required": True},
                {"name": "Objectifs", "required": True},
                {"name": "Périmètre", "required": True},
                {"name": "Exigences Fonctionnelles", "required": True},
                {"name": "Exigences Techniques", "required": True},
                {"name": "Architecture", "required": False},
                {"name": "Planning", "required": True},
                {"name": "Critères d'Acceptation", "required": True}
            ]
        }
    }
    
    def __init__(self):
        self.output_dir = "outputs/documents"
        os.makedirs(self.output_dir, exist_ok=True)
        self.generated_docs: Dict[str, GeneratedDocument] = {}
    
    async def generate(self, request: DocumentRequest) -> GeneratedDocument:
        """Génère un document DOCX"""
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
            from docx.oxml.ns import qn
        except ImportError:
            raise HTTPException(
                status_code=500,
                detail="python-docx non installé. Exécutez: pip install python-docx"
            )
        
        doc = Document()
        colors = self.COLORS.get(request.style, self.COLORS[DocumentStyle.PROFESSIONAL])
        
        # Configuration du document
        section = doc.sections[0]
        section.page_height = Inches(11.69)  # A4
        section.page_width = Inches(8.27)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        
        # En-tête
        if request.include_header:
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = f"{request.company} | {request.document_type.value.upper()}"
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Pied de page
        if request.include_footer:
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"© {datetime.now().year} {request.company} - Confidentiel"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Page de titre
        doc.add_paragraph()
        doc.add_paragraph()
        
        title = doc.add_heading(request.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if request.subtitle:
            subtitle = doc.add_paragraph(request.subtitle)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Métadonnées
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Préparé par: {request.author}\n")
        if request.recipient:
            meta.add_run(f"Pour: {request.recipient}\n")
            if request.recipient_company:
                meta.add_run(f"Entreprise: {request.recipient_company}\n")
        meta.add_run(f"Date: {datetime.now().strftime('%d %B %Y')}")
        
        # Saut de page
        doc.add_page_break()
        
        # Table des matières
        if request.include_toc:
            doc.add_heading("Table des Matières", level=1)
            doc.add_paragraph("(À mettre à jour dans Word: Clic droit → Mettre à jour le champ)")
            doc.add_page_break()
        
        # Sections du document
        for section_data in request.sections:
            section_title = section_data.get("title", "Section")
            section_content = section_data.get("content", "")
            section_level = section_data.get("level", 1)
            
            doc.add_heading(section_title, level=section_level)
            
            if isinstance(section_content, str):
                doc.add_paragraph(section_content)
            elif isinstance(section_content, list):
                for item in section_content:
                    if isinstance(item, str):
                        doc.add_paragraph(item, style="List Bullet")
                    elif isinstance(item, dict):
                        if item.get("type") == "paragraph":
                            doc.add_paragraph(item.get("text", ""))
                        elif item.get("type") == "bullet":
                            doc.add_paragraph(item.get("text", ""), style="List Bullet")
                        elif item.get("type") == "table":
                            self._add_table(doc, item.get("data", []))
        
        # Si pas de sections fournies, utiliser le template
        if not request.sections:
            template = self.TEMPLATES.get(request.document_type)
            if template:
                for section_info in template["sections"]:
                    doc.add_heading(section_info["name"], level=1)
                    doc.add_paragraph(f"[Contenu de la section {section_info['name']} à compléter]")
                    doc.add_paragraph()
        
        # Signature si contrat ou proposition
        if request.document_type in [DocumentType.CONTRACT, DocumentType.PROPOSAL]:
            doc.add_page_break()
            doc.add_heading("Acceptation", level=1)
            
            sig_table = doc.add_table(rows=2, cols=2)
            sig_table.style = "Table Grid"
            
            sig_table.cell(0, 0).text = f"{request.company}"
            sig_table.cell(0, 1).text = f"{request.recipient_company or 'Client'}"
            sig_table.cell(1, 0).text = f"\n\n\n{request.author}\nDate: ___/___/____"
            sig_table.cell(1, 1).text = f"\n\n\n{request.recipient or 'Représentant'}\nDate: ___/___/____"
        
        # Sauvegarder
        doc_id = f"doc_{datetime.now().strftime('%Y%m%d%H%M%S')}_{request.document_type.value}"
        filename = f"{doc_id}.docx"
        filepath = os.path.join(self.output_dir, filename)
        
        doc.save(filepath)
        
        # Créer l'objet résultat
        file_size = os.path.getsize(filepath)
        
        generated = GeneratedDocument(
            id=doc_id,
            filename=filename,
            filepath=filepath,
            document_type=request.document_type,
            created_at=datetime.now(),
            size_bytes=file_size,
            download_url=f"/skills/docx/download/{doc_id}"
        )
        
        self.generated_docs[doc_id] = generated
        
        return generated
    
    def _add_table(self, doc, data: List[List[str]]):
        """Ajoute un tableau au document"""
        if not data:
            return
        
        rows = len(data)
        cols = len(data[0]) if data else 0
        
        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        
        for i, row in enumerate(data):
            for j, cell_value in enumerate(row):
                table.cell(i, j).text = str(cell_value)
    
    def get_template(self, document_type: DocumentType) -> Dict[str, Any]:
        """Retourne le template pour un type de document"""
        return self.TEMPLATES.get(document_type, {"sections": []})
    
    def list_documents(self) -> List[GeneratedDocument]:
        """Liste tous les documents générés"""
        return list(self.generated_docs.values())


# Instance globale
docx_generator = DocxGenerator()


# Routes API

@router.post("/generate", response_model=Dict[str, Any])
async def generate_document(request: DocumentRequest):
    """
    Génère un document DOCX professionnel
    
    Types supportés:
    - proposal: Proposition commerciale
    - contract: Contrat
    - report: Rapport
    - invoice: Facture
    - letter: Lettre
    - specification: Cahier des charges
    """
    doc = await docx_generator.generate(request)
    return {
        "status": "success",
        "document_id": doc.id,
        "filename": doc.filename,
        "download_url": doc.download_url,
        "size_kb": round(doc.size_bytes / 1024, 2)
    }


@router.get("/templates/{document_type}")
async def get_template(document_type: DocumentType):
    """Retourne le template pour un type de document"""
    template = docx_generator.get_template(document_type)
    return {
        "document_type": document_type.value,
        "template": template,
        "required_sections": [s["name"] for s in template["sections"] if s.get("required")]
    }


@router.get("/templates")
async def list_templates():
    """Liste tous les templates disponibles"""
    return {
        doc_type.value: {
            "sections": template["sections"],
            "required_count": sum(1 for s in template["sections"] if s.get("required"))
        }
        for doc_type, template in docx_generator.TEMPLATES.items()
    }


@router.get("/documents")
async def list_documents():
    """Liste tous les documents générés"""
    return docx_generator.list_documents()


@router.get("/download/{doc_id}")
async def download_document(doc_id: str):
    """Télécharge un document généré"""
    from fastapi.responses import FileResponse
    
    doc = docx_generator.generated_docs.get(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return FileResponse(
        path=doc.filepath,
        filename=doc.filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@router.post("/from-prompt")
async def generate_from_prompt(
    prompt: str,
    document_type: DocumentType = DocumentType.PROPOSAL,
    style: DocumentStyle = DocumentStyle.PROFESSIONAL
):
    """
    Génère un document à partir d'un prompt en langage naturel
    Utilise Claude pour structurer le contenu
    """
    # TODO: Intégrer Claude API pour générer le contenu à partir du prompt
    
    return {
        "status": "pending",
        "message": "Cette fonctionnalité nécessite l'intégration Claude API",
        "prompt_received": prompt,
        "document_type": document_type.value
    }
