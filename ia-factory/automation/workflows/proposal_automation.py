"""
IA Factory Automation - Proposal Automation Workflow
Génération automatique de propositions commerciales
"""

from fastapi import APIRouter, HTTPException, BackgroundTasks
from pydantic import BaseModel, Field
from typing import Optional, List, Dict, Any
from datetime import datetime, timedelta
from enum import Enum
import json
import os

router = APIRouter(prefix="/proposals", tags=["Proposal Automation"])


class Market(str, Enum):
    CH = "CH"
    DZ = "DZ"


class Complexity(str, Enum):
    LOW = "Low"
    MEDIUM = "Medium"
    HIGH = "High"


class ServiceType(str, Enum):
    RAG = "RAG"
    MULTI_AGENT = "Multi-Agent"
    TRAINING = "Training"
    API_INTEGRATION = "API Integration"
    SUPPORT = "Support & Maintenance"
    VOICE_CLONING = "Voice Cloning"
    TEACHING_ASSISTANT = "Teaching Assistant"


class ProposalRequest(BaseModel):
    """Requête de génération de proposition"""
    client_name: str
    company_name: str
    company_industry: str
    contact_email: str
    market: Market
    services: List[ServiceType]
    scope: Dict[str, Any] = Field(
        default_factory=dict,
        description="Ex: {'num_documents': 10000, 'num_users': 50}"
    )
    complexity: Complexity = Complexity.MEDIUM
    timeline_months: int = Field(default=3, ge=1, le=24)
    notes: Optional[str] = None


class ProposalTier(BaseModel):
    """Tier de proposition"""
    name: str
    price: float
    currency: str
    features: List[str]
    recommended: bool = False


class Proposal(BaseModel):
    """Proposition générée"""
    id: str
    client_name: str
    company_name: str
    market: Market
    created_at: datetime
    valid_until: datetime
    services: List[ServiceType]
    pricing_breakdown: List[Dict[str, Any]]
    tiers: List[ProposalTier]
    total_price: float
    currency: str
    timeline_weeks: int
    status: str = "draft"
    docx_path: Optional[str] = None
    pdf_path: Optional[str] = None


class PricingCalculator:
    """Calculateur de prix intelligent"""
    
    # Prix de base par service (CHF)
    BASE_PRICES_CH = {
        ServiceType.RAG: 25000,
        ServiceType.MULTI_AGENT: 40000,
        ServiceType.TRAINING: 8000,
        ServiceType.API_INTEGRATION: 5000,
        ServiceType.SUPPORT: 2000,
        ServiceType.VOICE_CLONING: 5000,
        ServiceType.TEACHING_ASSISTANT: 15000,
    }
    
    # Prix de base par service (DZD)
    BASE_PRICES_DZ = {
        ServiceType.RAG: 800000,
        ServiceType.MULTI_AGENT: 1500000,
        ServiceType.TRAINING: 300000,
        ServiceType.API_INTEGRATION: 200000,
        ServiceType.SUPPORT: 80000,
        ServiceType.VOICE_CLONING: 200000,
        ServiceType.TEACHING_ASSISTANT: 500000,
    }
    
    # Multiplicateurs complexité
    COMPLEXITY_MULTIPLIERS = {
        Complexity.LOW: 1.0,
        Complexity.MEDIUM: 1.5,
        Complexity.HIGH: 2.5,
    }
    
    def calculate(
        self,
        services: List[ServiceType],
        scope: Dict[str, Any],
        complexity: Complexity,
        market: Market,
        timeline_months: int
    ) -> Dict[str, Any]:
        """Calcule le pricing complet"""
        
        base_prices = self.BASE_PRICES_CH if market == Market.CH else self.BASE_PRICES_DZ
        currency = "CHF" if market == Market.CH else "DZD"
        multiplier = self.COMPLEXITY_MULTIPLIERS[complexity]
        
        breakdown = []
        total = 0
        
        for service in services:
            base = base_prices[service]
            adjusted = base * multiplier
            
            # Ajustements scope
            if service == ServiceType.RAG:
                num_docs = scope.get("num_documents", 10000)
                if num_docs > 50000:
                    adjusted *= 1.5
                elif num_docs > 20000:
                    adjusted *= 1.2
            
            if service == ServiceType.MULTI_AGENT:
                num_agents = scope.get("num_agents", 3)
                adjusted *= (1 + (num_agents - 1) * 0.15)
            
            if service == ServiceType.TEACHING_ASSISTANT:
                num_subjects = scope.get("num_subjects", 3)
                adjusted *= (1 + (num_subjects - 1) * 0.2)
            
            total += adjusted
            
            breakdown.append({
                "service": service.value,
                "base_price": base,
                "complexity_multiplier": multiplier,
                "scope_adjustments": "Applied",
                "final_price": round(adjusted, 2)
            })
        
        # Support mensuel si durée > 3 mois
        monthly_support = 0
        if timeline_months > 3:
            support_rate = base_prices[ServiceType.SUPPORT]
            monthly_support = support_rate * (timeline_months - 3)
            breakdown.append({
                "service": "Support Continu",
                "base_price": support_rate,
                "months": timeline_months - 3,
                "final_price": monthly_support
            })
            total += monthly_support
        
        # Générer 3 tiers
        tiers = self._generate_tiers(total, currency, services)
        
        return {
            "breakdown": breakdown,
            "total": round(total, 2),
            "currency": currency,
            "tiers": tiers,
            "timeline_weeks": timeline_months * 4
        }
    
    def _generate_tiers(
        self,
        base_total: float,
        currency: str,
        services: List[ServiceType]
    ) -> List[ProposalTier]:
        """Génère les 3 tiers de pricing"""
        
        essential_features = [
            "Core features",
            "Email support",
            "Documentation",
            "3 mois garantie"
        ]
        
        pro_features = essential_features + [
            "Toutes les features",
            "Support prioritaire",
            "Formation équipe incluse",
            "SLA 99% uptime",
            "Dashboard analytics"
        ]
        
        enterprise_features = pro_features + [
            "Features custom",
            "Support dédié 24/7",
            "On-premise option",
            "SLA 99.9% uptime",
            "Training certifiant",
            "Maintenance incluse 12 mois"
        ]
        
        return [
            ProposalTier(
                name="Essential",
                price=round(base_total * 0.7, 2),
                currency=currency,
                features=essential_features,
                recommended=False
            ),
            ProposalTier(
                name="Professional",
                price=round(base_total, 2),
                currency=currency,
                features=pro_features,
                recommended=True  # Recommandé
            ),
            ProposalTier(
                name="Enterprise",
                price=round(base_total * 1.5, 2),
                currency=currency,
                features=enterprise_features,
                recommended=False
            )
        ]


class ProposalGenerator:
    """Générateur de propositions"""
    
    def __init__(self):
        self.pricing = PricingCalculator()
        self.proposals_db: Dict[str, Proposal] = {}
    
    async def generate_proposal(self, request: ProposalRequest) -> Proposal:
        """Génère une proposition complète"""
        
        # Calculer pricing
        pricing_data = self.pricing.calculate(
            services=request.services,
            scope=request.scope,
            complexity=request.complexity,
            market=request.market,
            timeline_months=request.timeline_months
        )
        
        # Créer ID unique
        proposal_id = f"prop_{datetime.now().strftime('%Y%m%d%H%M%S')}_{request.company_name[:10].replace(' ', '_')}"
        
        # Créer proposition
        proposal = Proposal(
            id=proposal_id,
            client_name=request.client_name,
            company_name=request.company_name,
            market=request.market,
            created_at=datetime.now(),
            valid_until=datetime.now() + timedelta(days=30),
            services=request.services,
            pricing_breakdown=pricing_data["breakdown"],
            tiers=[ProposalTier(**t.model_dump()) for t in pricing_data["tiers"]],
            total_price=pricing_data["total"],
            currency=pricing_data["currency"],
            timeline_weeks=pricing_data["timeline_weeks"],
            status="draft"
        )
        
        # Sauvegarder
        self.proposals_db[proposal_id] = proposal
        
        return proposal
    
    async def generate_docx(self, proposal_id: str) -> str:
        """Génère le document DOCX de la proposition"""
        
        if proposal_id not in self.proposals_db:
            raise ValueError("Proposal not found")
        
        proposal = self.proposals_db[proposal_id]
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Style titre
            title = doc.add_heading("PROPOSITION COMMERCIALE", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Sous-titre
            subtitle = doc.add_paragraph(f"IA Factory × {proposal.company_name}")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date
            date_para = doc.add_paragraph(f"Date: {proposal.created_at.strftime('%d/%m/%Y')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Espace
            
            # Section Client
            doc.add_heading("1. Client", level=1)
            doc.add_paragraph(f"Entreprise: {proposal.company_name}")
            doc.add_paragraph(f"Contact: {proposal.client_name}")
            doc.add_paragraph(f"Marché: {'Suisse' if proposal.market == Market.CH else 'Algérie'}")
            
            # Section Services
            doc.add_heading("2. Services Proposés", level=1)
            for service in proposal.services:
                doc.add_paragraph(f"• {service.value}", style="List Bullet")
            
            # Section Pricing
            doc.add_heading("3. Investissement", level=1)
            
            # Table pricing
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            
            # Header
            header_cells = table.rows[0].cells
            header_cells[0].text = "Tier"
            header_cells[1].text = "Prix"
            header_cells[2].text = "Recommandé"
            
            # Rows
            for tier in proposal.tiers:
                row = table.add_row().cells
                row[0].text = tier.name
                row[1].text = f"{tier.price:,.0f} {tier.currency}"
                row[2].text = "✓" if tier.recommended else ""
            
            # Section Timeline
            doc.add_heading("4. Timeline", level=1)
            doc.add_paragraph(f"Durée estimée: {proposal.timeline_weeks} semaines")
            
            # Section Validité
            doc.add_heading("5. Validité", level=1)
            doc.add_paragraph(f"Cette proposition est valable jusqu'au {proposal.valid_until.strftime('%d/%m/%Y')}")
            
            # Signature
            doc.add_paragraph()
            doc.add_paragraph()
            signature = doc.add_paragraph("Boualem Chebaki")
            signature.add_run("\nFondateur, IA Factory")
            signature.add_run("\ncontact@iafactory.ch | www.iafactory.ch")
            
            # Sauvegarder
            output_dir = "outputs/proposals"
            os.makedirs(output_dir, exist_ok=True)
            output_path = f"{output_dir}/{proposal_id}.docx"
            doc.save(output_path)
            
            # Update proposal
            proposal.docx_path = output_path
            proposal.status = "generated"
            
            return output_path
            
        except ImportError:
            # Fallback si python-docx pas installé
            return "python-docx required"
    
    def get_proposal(self, proposal_id: str) -> Optional[Proposal]:
        """Récupère une proposition"""
        return self.proposals_db.get(proposal_id)
    
    def list_proposals(self) -> List[Proposal]:
        """Liste toutes les propositions"""
        return list(self.proposals_db.values())


# Instance globale
proposal_generator = ProposalGenerator()


# Routes API

@router.post("/generate", response_model=Dict[str, Any])
async def generate_proposal(request: ProposalRequest, background_tasks: BackgroundTasks):
    """
    Génère une nouvelle proposition commerciale
    """
    proposal = await proposal_generator.generate_proposal(request)
    
    # Générer DOCX en background
    background_tasks.add_task(proposal_generator.generate_docx, proposal.id)
    
    return {
        "status": "success",
        "proposal_id": proposal.id,
        "total_price": proposal.total_price,
        "currency": proposal.currency,
        "recommended_tier": next(
            (t.name for t in proposal.tiers if t.recommended),
            "Professional"
        ),
        "valid_until": proposal.valid_until.isoformat(),
        "message": "Proposition générée. Document DOCX en cours de création."
    }


@router.get("/calculate-price")
async def calculate_price(
    services: str,  # comma-separated
    market: Market,
    complexity: Complexity = Complexity.MEDIUM,
    timeline_months: int = 3
):
    """
    Calcule le prix sans créer de proposition
    """
    service_list = [ServiceType(s.strip()) for s in services.split(",")]
    
    pricing = PricingCalculator().calculate(
        services=service_list,
        scope={},
        complexity=complexity,
        market=market,
        timeline_months=timeline_months
    )
    
    return {
        "services": [s.value for s in service_list],
        "market": market.value,
        "complexity": complexity.value,
        "breakdown": pricing["breakdown"],
        "total": pricing["total"],
        "currency": pricing["currency"],
        "tiers": [
            {
                "name": t.name,
                "price": t.price,
                "recommended": t.recommended
            }
            for t in pricing["tiers"]
        ]
    }


@router.get("/{proposal_id}")
async def get_proposal(proposal_id: str):
    """Récupère une proposition par ID"""
    proposal = proposal_generator.get_proposal(proposal_id)
    
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")
    
    return proposal


@router.get("/{proposal_id}/download")
async def download_proposal(proposal_id: str, format: str = "docx"):
    """Télécharge la proposition en DOCX ou PDF"""
    proposal = proposal_generator.get_proposal(proposal_id)
    
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")
    
    if format == "docx" and proposal.docx_path:
        return {"file_path": proposal.docx_path}
    
    raise HTTPException(status_code=400, detail="Document not ready or format not supported")


@router.get("/")
async def list_proposals(status: Optional[str] = None, market: Optional[Market] = None):
    """Liste toutes les propositions"""
    proposals = proposal_generator.list_proposals()
    
    if status:
        proposals = [p for p in proposals if p.status == status]
    
    if market:
        proposals = [p for p in proposals if p.market == market]
    
    return proposals
